from dataclasses import dataclass
from io import BytesIO
from pathlib import PurePath
from uuid import UUID

import fitz
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy import bindparam, func, insert
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.models.schemas import DocumentIngestResponse
from app.models.tables import DocumentChunk
from app.services.embedding import HuggingFaceEmbeddingClient


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int
    text: str


@dataclass(frozen=True)
class ChunkPayload:
    document_name: str
    page_number: int
    parent_content: str
    child_content: str


ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}


async def ingest_document(
    *,
    db: AsyncSession,
    user_id: UUID,
    upload: UploadFile,
    embedding_client: HuggingFaceEmbeddingClient | None = None,
) -> DocumentIngestResponse:
    settings = get_settings()
    document_name = _safe_document_name(upload.filename)
    extension = PurePath(document_name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Only .txt, .pdf, and .docx documents are supported.",
        )

    raw_bytes = await _read_upload_limited(upload, settings.max_upload_bytes)
    try:
        pages = _extract_pages(raw_bytes, extension)
        chunks = _build_parent_child_chunks(document_name, pages)
        if not chunks:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="No readable text was found in the document.",
            )

        client = embedding_client or HuggingFaceEmbeddingClient()
        embeddings = await client.embed_texts([chunk.child_content for chunk in chunks])
        await _insert_chunks(db=db, user_id=user_id, chunks=chunks, embeddings=embeddings)

        parent_count = len({(chunk.page_number, chunk.parent_content) for chunk in chunks})
        return DocumentIngestResponse(
            document_name=document_name,
            pages_processed=len(pages),
            parent_chunks=parent_count,
            child_chunks=len(chunks),
        )
    finally:
        # Drop references promptly; the process-and-discard pipeline keeps persistence in Postgres only.
        raw_bytes = b""
        await upload.close()


async def _read_upload_limited(upload: UploadFile, max_bytes: int) -> bytes:
    payload = await upload.read(max_bytes + 1)
    if len(payload) > max_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"Upload exceeds the {max_bytes // (1024 * 1024)}MB limit.",
        )
    return payload


def _safe_document_name(filename: str | None) -> str:
    name = PurePath(filename or "document.txt").name.strip()
    return name or "document.txt"


def _extract_pages(raw_bytes: bytes, extension: str) -> list[ExtractedPage]:
    if extension == ".pdf":
        return _extract_pdf_pages(raw_bytes)
    if extension == ".docx":
        return _extract_docx_pages(raw_bytes)
    return _extract_text_pages(raw_bytes)


def _extract_pdf_pages(raw_bytes: bytes) -> list[ExtractedPage]:
    document = fitz.open(stream=raw_bytes, filetype="pdf")
    try:
        pages = [
            ExtractedPage(page_number=index + 1, text=page.get_text("text").strip())
            for index, page in enumerate(document)
        ]
    finally:
        document.close()
    return [page for page in pages if page.text]


def _extract_docx_pages(raw_bytes: bytes) -> list[ExtractedPage]:
    document = DocxDocument(BytesIO(raw_bytes))
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    text = "\n\n".join(blocks).strip()
    return [ExtractedPage(page_number=1, text=text)] if text else []


def _extract_text_pages(raw_bytes: bytes) -> list[ExtractedPage]:
    text = _decode_text(raw_bytes)
    page_texts = [page.strip() for page in text.split("\f")]
    return [
        ExtractedPage(page_number=index + 1, text=page_text)
        for index, page_text in enumerate(page_texts)
        if page_text
    ]


def _decode_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def _build_parent_child_chunks(document_name: str, pages: list[ExtractedPage]) -> list[ChunkPayload]:
    parent_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    child_splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=50)

    chunks: list[ChunkPayload] = []
    for page in pages:
        parent_chunks = parent_splitter.split_text(page.text)
        for parent in parent_chunks:
            parent = parent.strip()
            if not parent:
                continue
            for child in child_splitter.split_text(parent):
                child = child.strip()
                if child:
                    chunks.append(
                        ChunkPayload(
                            document_name=document_name,
                            page_number=page.page_number,
                            parent_content=parent,
                            child_content=child,
                        )
                    )
    return chunks


async def _insert_chunks(
    *,
    db: AsyncSession,
    user_id: UUID,
    chunks: list[ChunkPayload],
    embeddings: list[list[float]],
) -> None:
    if len(chunks) != len(embeddings):
        raise RuntimeError("Embedding count does not match chunk count.")

    rows = [
        {
            "user_id": user_id,
            "document_name": chunk.document_name,
            "page_number": chunk.page_number,
            "child_content": chunk.child_content,
            "parent_content": chunk.parent_content,
            "embedding": embedding,
            "search_text": f"{chunk.child_content}\n\n{chunk.parent_content}",
        }
        for chunk, embedding in zip(chunks, embeddings, strict=True)
    ]

    statement = insert(DocumentChunk).values(
        user_id=bindparam("user_id"),
        document_name=bindparam("document_name"),
        page_number=bindparam("page_number"),
        child_content=bindparam("child_content"),
        parent_content=bindparam("parent_content"),
        embedding=bindparam("embedding"),
        text_search_vector=func.to_tsvector("english", bindparam("search_text")),
    )

    await db.execute(statement, rows)
    await db.commit()
